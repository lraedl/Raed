"""
ЛР №7 — Элементарные частицы (ООП-версия ЛР №6)
GUI: FreeSimpleGUI (PySimpleGUI-совместимый бесплатный форк)

Иерархия классов:
    Particle (ABC) — абстрактный базовый класс
        ├── ChargedParticle   — заряженная частица (электрон, протон)
        └── NeutralParticle   — нейтральная частица (нейтрон)
"""

from abc import ABC, abstractmethod
import FreeSimpleGUI as sg
from docx import Document
import openpyxl

# ══════════════════════════════════════════════════════
#  Физические константы
# ══════════════════════════════════════════════════════
H = 6.626e-34   # постоянная Планка (Дж·с)
C = 3e8         # скорость света (м/с)


# ══════════════════════════════════════════════════════
#  Абстрактный базовый класс
# ══════════════════════════════════════════════════════
class Particle(ABC):
    """
    Абстрактный класс — шаблон для любой элементарной частицы.
    Нельзя создать объект этого класса напрямую.
    """

    def __init__(self, name: str, mass: float):
        self.name = name    # имя частицы (через property ниже)
        self.mass = mass    # масса (через property ниже)

    # ── managed-атрибут: name ─────────────────────────
    @property
    def name(self):
        return self._name

    @name.setter
    def name(self, value):
        if not isinstance(value, str) or not value:
            raise ValueError("Имя должно быть непустой строкой")
        self._name = value

    # ── managed-атрибут: mass ─────────────────────────
    @property
    def mass(self):
        return self._mass

    @mass.setter
    def mass(self, value):
        if value <= 0:
            raise ValueError("Масса должна быть больше нуля")
        self._mass = value

    # ── Абстрактные методы — каждый подкласс ОБЯЗАН их реализовать ──
    @abstractmethod
    def udelny_zaryad(self) -> float:
        """Удельный заряд (Кл/кг)."""

    @abstractmethod
    def kompton(self) -> float:
        """Комптоновская длина волны (м)."""

    @abstractmethod
    def info(self) -> dict:
        """Словарь с результатами расчётов для отчёта."""

    # ── dunder-методы ─────────────────────────────────
    def __str__(self):
        # Вызывается при print(particle)
        return f"{self.name}: масса={self.mass:.3e} кг"

    def __repr__(self):
        # Вызывается при выводе в интерпретаторе
        return f"Particle(name={self._name!r}, mass={self._mass!r})"


# ══════════════════════════════════════════════════════
#  Подкласс 1: заряженная частица (электрон, протон)
# ══════════════════════════════════════════════════════
class ChargedParticle(Particle):
    """Частица с ненулевым зарядом."""

    def __init__(self, name: str, mass: float, charge: float):
        super().__init__(name, mass)   # вызываем __init__ родителя
        self.charge = charge           # заряд (через property)

    # ── managed-атрибут: charge ───────────────────────
    @property
    def charge(self):
        return self._charge

    @charge.setter
    def charge(self, value):
        if value <= 0:
            raise ValueError("Заряд заряженной частицы должен быть > 0")
        self._charge = value

    # ── Реализация абстрактных методов ────────────────
    def udelny_zaryad(self) -> float:
        return self._charge / self._mass

    def kompton(self) -> float:
        return H / (self._mass * C)

    def info(self) -> dict:
        return {
            "name":        self._name,
            "mass":        self._mass,
            "charge":      self._charge,
            "spec_charge": self.udelny_zaryad(),
            "compton":     self.kompton(),
        }

    # ── dunder-методы ─────────────────────────────────
    def __str__(self):
        return (f"{self._name}: масса={self._mass:.3e} кг, "
                f"заряд={self._charge:.3e} Кл")

    def __eq__(self, other):
        # Две частицы равны, если у них одинаковые масса и заряд
        if not isinstance(other, ChargedParticle):
            return False
        return self._mass == other._mass and self._charge == other._charge


# ══════════════════════════════════════════════════════
#  Подкласс 2: нейтральная частица (нейтрон)
# ══════════════════════════════════════════════════════
class NeutralParticle(Particle):
    """Частица без заряда — нейтрон."""

    def __init__(self, name: str, mass: float):
        super().__init__(name, mass)

    # ── Реализация абстрактных методов ────────────────
    def udelny_zaryad(self) -> float:
        return 0.0   # нейтрон не имеет заряда

    def kompton(self) -> float:
        return H / (self._mass * C)

    def info(self) -> dict:
        return {
            "name":        self._name,
            "mass":        self._mass,
            "charge":      0.0,
            "spec_charge": 0.0,
            "compton":     self.kompton(),
        }

    # ── dunder-методы ─────────────────────────────────
    def __str__(self):
        return f"{self._name}: масса={self._mass:.3e} кг, заряд=0 (нейтральная)"

    def __eq__(self, other):
        if not isinstance(other, NeutralParticle):
            return False
        return self._mass == other._mass


# ══════════════════════════════════════════════════════
#  Класс для сохранения отчётов
# ══════════════════════════════════════════════════════
class Report:
    """Сохраняет список результатов в Word или Excel."""

    def __init__(self, results: list):
        # results — список словарей из метода info() каждой частицы
        self.results = results

    def __len__(self):
        # len(report) вернёт количество частиц в отчёте
        return len(self.results)

    def __repr__(self):
        return f"Report({len(self.results)} частиц)"

    def v_word(self, path: str):
        doc = Document()
        doc.add_heading("Элементарные частицы", level=1)
        for r in self.results:
            doc.add_heading(r["name"], level=2)
            doc.add_paragraph(f"Масса:                    {r['mass']:.3e} кг")
            doc.add_paragraph(f"Заряд:                    {r['charge']:.3e} Кл")
            doc.add_paragraph(f"Удельный заряд:           {r['spec_charge']:.3e} Кл/кг")
            doc.add_paragraph(f"Комптоновская длина волны:{r['compton']:.3e} м")
        doc.save(path)

    def v_excel(self, path: str):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["Частица", "Масса (кг)", "Заряд (Кл)",
                   "Удельный заряд (Кл/кг)", "Комптон (м)"])
        for r in self.results:
            ws.append([r["name"], r["mass"], r["charge"],
                       r["spec_charge"], r["compton"]])
        wb.save(path)


# ══════════════════════════════════════════════════════
#  Список всех частиц
# ══════════════════════════════════════════════════════
PARTICLES = [
    ChargedParticle("Электрон", mass=9.109e-31,  charge=1.602e-19),
    ChargedParticle("Протон",   mass=1.673e-27,  charge=1.602e-19),
    NeutralParticle("Нейтрон",  mass=1.675e-27),
]

# Словарь для быстрого поиска по имени
PARTICLES_BY_NAME = {p.name: p for p in PARTICLES}


# ══════════════════════════════════════════════════════
#  GUI на FreeSimpleGUI
# ══════════════════════════════════════════════════════

sg.theme("LightBlue2")   # тема оформления

# Описываем разметку окна — список строк с виджетами
layout = [
    [sg.Text("Элементарные частицы", font=("Arial", 14, "bold"))],
    [sg.HorizontalSeparator()],

    # Строка выбора частицы
    [sg.Text("Выберите частицу:"),
     sg.Combo(list(PARTICLES_BY_NAME.keys()), default_value="Электрон",
              key="-COMBO-", readonly=True, size=(15, 1))],

    [sg.Button("Рассчитать", button_color=("white", "#4CAF50"))],

    [sg.HorizontalSeparator()],

    # Поля результатов (readonly — только для чтения)
    [sg.Text("Масса:"),                    sg.Input("—", key="-MASS-",    size=(25,1), readonly=True)],
    [sg.Text("Заряд:"),                    sg.Input("—", key="-CHARGE-",  size=(25,1), readonly=True)],
    [sg.Text("Удельный заряд:"),           sg.Input("—", key="-SPEC-",    size=(25,1), readonly=True)],
    [sg.Text("Комптоновская длина волны:"), sg.Input("—", key="-COMPTON-", size=(25,1), readonly=True)],

    [sg.HorizontalSeparator()],

    # Кнопки сохранения текущей частицы
    [sg.Text("Сохранить текущую частицу:")],
    [sg.Button("💾 Word (.docx)", key="-WORD1-"),
     sg.Button("💾 Excel (.xlsx)", key="-XLSX1-")],

    # Кнопки сохранения всех частиц
    [sg.Text("Сохранить все три частицы:")],
    [sg.Button("📄 Все в Word", key="-WORD_ALL-"),
     sg.Button("📊 Все в Excel", key="-XLSX_ALL-")],
]

window = sg.Window("Элементарные частицы — ЛР №7", layout)

current_result = None   # здесь будем хранить последний расчёт

# Главный цикл — обрабатываем события окна
while True:
    event, values = window.read()

    # Закрытие окна
    if event == sg.WIN_CLOSED:
        break

    # ── Кнопка "Рассчитать" ───────────────────────────
    if event == "Рассчитать":
        name = values["-COMBO-"]
        particle = PARTICLES_BY_NAME[name]   # получаем объект частицы

        # Вызываем методы объекта
        spec    = particle.udelny_zaryad()
        compton = particle.kompton()
        data    = particle.info()

        # Обновляем поля в окне
        window["-MASS-"].update(   f"{data['mass']:.3e} кг")
        window["-CHARGE-"].update( f"{data['charge']:.3e} Кл")
        window["-SPEC-"].update(   f"{data['spec_charge']:.3e} Кл/кг")
        window["-COMPTON-"].update(f"{data['compton']:.3e} м")

        current_result = data   # запоминаем для сохранения

    # ── Сохранение текущей частицы в Word ────────────
    if event == "-WORD1-":
        if current_result is None:
            sg.popup_error("Сначала нажмите «Рассчитать»!")
        else:
            path = sg.popup_get_file("Сохранить как", save_as=True,
                                     file_types=(("Word", "*.docx"),),
                                     default_extension=".docx")
            if path:
                Report([current_result]).v_word(path)
                sg.popup(f"Сохранено:\n{path}")

    # ── Сохранение текущей частицы в Excel ───────────
    if event == "-XLSX1-":
        if current_result is None:
            sg.popup_error("Сначала нажмите «Рассчитать»!")
        else:
            path = sg.popup_get_file("Сохранить как", save_as=True,
                                     file_types=(("Excel", "*.xlsx"),),
                                     default_extension=".xlsx")
            if path:
                Report([current_result]).v_excel(path)
                sg.popup(f"Сохранено:\n{path}")

    # ── Сохранение всех частиц в Word ─────────────────
    if event == "-WORD_ALL-":
        path = sg.popup_get_file("Сохранить как", save_as=True,
                                 file_types=(("Word", "*.docx"),),
                                 default_extension=".docx")
        if path:
            all_results = [p.info() for p in PARTICLES]  # info() у каждого объекта
            Report(all_results).v_word(path)
            sg.popup(f"Сохранено:\n{path}")

    # ── Сохранение всех частиц в Excel ────────────────
    if event == "-XLSX_ALL-":
        path = sg.popup_get_file("Сохранить как", save_as=True,
                                 file_types=(("Excel", "*.xlsx"),),
                                 default_extension=".xlsx")
        if path:
            all_results = [p.info() for p in PARTICLES]
            Report(all_results).v_excel(path)
            sg.popup(f"Сохранено:\n{path}")

window.close()
