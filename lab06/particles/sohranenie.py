# sohranenie.py — сохранение результатов в Word и Excel

from docx import Document
import openpyxl


def v_word(results, path):
    # Создаём новый Word-документ
    doc = Document()
    doc.add_heading("Элементарные частицы", level=1)

    for r in results:
        doc.add_heading(r["name"], level=2)
        doc.add_paragraph(f"Масса:                    {r['mass']:.3e} кг")
        doc.add_paragraph(f"Заряд:                    {r['charge']:.3e} Кл")
        doc.add_paragraph(f"Удельный заряд:           {r['spec_charge']:.3e} Кл/кг")
        doc.add_paragraph(f"Комптоновская длина волны:{r['compton']:.3e} м")

    doc.save(path)  # сохраняем файл


def v_excel(results, path):
    # Создаём новую Excel-таблицу
    wb = openpyxl.Workbook()
    ws = wb.active

    # Первая строка — заголовки
    ws.append(["Частица", "Масса (кг)", "Заряд (Кл)", "Удельный заряд (Кл/кг)", "Комптон (м)"])

    # Остальные строки — данные
    for r in results:
        ws.append([r["name"], r["mass"], r["charge"], r["spec_charge"], r["compton"]])

    wb.save(path)  # сохраняем файл
